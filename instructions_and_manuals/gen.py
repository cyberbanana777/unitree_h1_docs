from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_vmware_instruction_docx():
    # Создаем документ
    doc = Document()
    
    # Настраиваем стили для русского языка и шрифтов
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Устанавливаем межстрочный интервал 1.5
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    
    # Добавляем заголовок
    title = doc.add_heading('🖥️ Как установить VMware', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Добавляем подзаголовок курсивом
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Пошаговая инструкция для новичков')
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Добавляем разделитель
    doc.add_paragraph('―' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Шаг 1
    doc.add_heading('🔽 Шаг 1. Скачиваем VMware', level=2)
    
    p1 = doc.add_paragraph('Перейдите по ссылке ниже, чтобы скачать установочный файл программы VMware:')
    
    # Добавляем гиперссылку
    p_link = doc.add_paragraph()
    p_link.add_run('🔗 ')
    link_run = p_link.add_run('Скачать VMware')
    link_run.font.color.rgb = (0, 0, 255)  # Синий цвет
    # Для реальной гиперссылки нужно использовать add_hyperlink, но это сложнее
    
    p2 = doc.add_paragraph('После нажатия на кнопку "скачать" начнётся загрузка. Дождитесь её завершения.')
    
    # Добавляем изображение 1
    try:
        doc.add_picture('photo_for_instructions/VMWare/dowload1.png', width=Inches(5))
        p_img1 = doc.add_paragraph('Куда нажать, чтобы скачать VMware')
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.style.font.size = Pt(12)
    except:
        doc.add_paragraph('[Изображение: Куда нажать, чтобы скачать VMware]')
    
    p3 = doc.add_paragraph('Запустите скачанный файл.')
    
    # Добавляем изображение 2
    try:
        doc.add_picture('photo_for_instructions/VMWare/instal2.png', width=Inches(5))
        p_img2 = doc.add_paragraph('Куда нажать, чтобы установить VMware')
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.style.font.size = Pt(12)
    except:
        doc.add_paragraph('[Изображение: Куда нажать, чтобы установить VMware]')
    
    doc.add_paragraph('―' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Шаг 2
    doc.add_heading('⚙️ Шаг 2. Устанавливаем VMware', level=2)
    
    p4 = doc.add_paragraph('Установка проходит в несколько простых шагов — просто следуйте подсказкам на экране. Ниже показан весь процесс по порядку:')
    
    # Добавляем все изображения установки
    setup_images = [
        ('setup1.PNG', 'Установочное окно №1'),
        ('setup2.PNG', 'Принять лицензионное соглашение'),
        ('setup3.PNG', '1 - пункт - элективный (Enhanced Keyboard Driver)\n2 - пункт - обязательный (Add VWware Workstation tools into PATH)'),
        ('setup4.PNG', 'Выбор следующих пунктов - по желанию'),
        ('setup5.PNG', 'Выбор следующих пунктов - по желанию'),
        ('setup6.PNG', 'Установочное окно №6'),
        ('setup7.PNG', 'Установочное окно №7')
    ]
    
    for img_file, description in setup_images:
        try:
            doc.add_picture(f'photo_for_instructions/VMWare/{img_file}', width=Inches(5))
            p_desc = doc.add_paragraph(description)
            p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_desc.style.font.size = Pt(12)
        except:
            doc.add_paragraph(f'[Изображение: {description}]')
    
    doc.add_paragraph('―' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Шаг 3
    doc.add_heading('✅ Шаг 3. Подтверждаем установку сетевых компонентов', level=2)
    
    p5 = doc.add_paragraph('Если во время установки появится окно с вопросом о разрешении перезагрузки. Нажмите «Yes».')
    
    try:
        doc.add_picture('photo_for_instructions/VMWare/setup8.PNG', width=Inches(5))
        p_img8 = doc.add_paragraph('Установочное окно №8')
        p_img8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img8.style.font.size = Pt(12)
    except:
        doc.add_paragraph('[Изображение: Установочное окно №8]')
    
    doc.add_paragraph('―' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Шаг 4
    doc.add_heading('🖱️ Шаг 4. Запускаем VMware', level=2)
    
    p6 = doc.add_paragraph('После завершения установки на рабочем столе появится ярлык программы. Дважды кликните по нему, чтобы запустить VMware.')
    
    try:
        doc.add_picture('photo_for_instructions/VMWare/icon.PNG', width=Inches(3))
        p_icon = doc.add_paragraph('Иконка VMware')
        p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_icon.style.font.size = Pt(12)
    except:
        doc.add_paragraph('[Изображение: Иконка VMware]')
    
    p7 = doc.add_paragraph('Затем выберите настройки, как показано на изображении ниже:')
    
    open_images = [
        ('open1.PNG', 'Выбор лицензии использования'),
        ('open2.PNG', 'Приветственное окно')
    ]
    
    for img_file, description in open_images:
        try:
            doc.add_picture(f'photo_for_instructions/VMWare/{img_file}', width=Inches(5))
            p_open = doc.add_paragraph(description)
            p_open.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_open.style.font.size = Pt(12)
        except:
            doc.add_paragraph(f'[Изображение: {description}]')
    
    doc.add_paragraph('―' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Завершающий раздел
    doc.add_heading('🎯 Готово! Вот как выглядит главный интерфейс VMware', level=2)
    
    p8 = doc.add_paragraph('Теперь вы видите основное окно программы. Здесь вы сможете создавать, запускать и управлять своими виртуальными машинами.')
    
    try:
        doc.add_picture('photo_for_instructions/VMWare/interface.PNG', width=Inches(6))
        p_interface = doc.add_paragraph('Интерфейс VMware')
        p_interface.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_interface.style.font.size = Pt(12)
    except:
        doc.add_paragraph('[Изображение: Интерфейс VMware]')
    
    # Сохраняем документ
    doc.save('VMware_installation_instructions.docx')
    print("DOCX-файл успешно создан: VMware_installation_instructions.docx")

if __name__ == "__main__":
    create_vmware_instruction_docx()